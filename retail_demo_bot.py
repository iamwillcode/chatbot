from PyQt5.QtWidgets import QMainWindow, QTabWidget, QWidget, QVBoxLayout, QHBoxLayout, QLineEdit, QPushButton, QCheckBox, QComboBox, QTextEdit, QScrollArea, QLabel, QDialog, QSlider, QFileDialog, QProgressBar, QMessageBox, QTreeView, QMenu, QToolBar, QAction
from PyQt5.QtGui import QFont, QPixmap, QStandardItemModel, QStandardItem
from PyQt5.QtCore import Qt
from tinydb import TinyDB, Query
from fuzzywuzzy import fuzz
from docx import Document
import pdfplumber
import fitz  # PyMuPDF
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from collections import Counter
import json
import os
from pathlib import Path
import re
import uuid
import logging
from PIL import Image

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Download NLTK data
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)

class ImageZoomDialog(QDialog):
    def __init__(self, image_path, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Image Preview")
        self.setGeometry(100, 100, 600, 400)
        self.pixmap = QPixmap(image_path)
        self.scale = 1.0
        self.offset = [0, 0]
        self.last_pos = None

        layout = QVBoxLayout(self)
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.image_label = QLabel()
        self.image_label.setAlignment(Qt.AlignCenter)
        self.scroll_area.setWidget(self.image_label)
        layout.addWidget(self.scroll_area)

        zoom_layout = QHBoxLayout()
        self.zoom_slider = QSlider(Qt.Horizontal)
        self.zoom_slider.setRange(50, 200)
        self.zoom_slider.setValue(100)
        self.zoom_slider.valueChanged.connect(self.update_zoom)
        zoom_layout.addWidget(QLabel("Zoom:"))
        zoom_layout.addWidget(self.zoom_slider)
        layout.addLayout(zoom_layout)

        self.image_label.setMouseTracking(True)
        self.image_label.installEventFilter(self)
        self.update_image()

    def update_image(self):
        scaled_pixmap = self.pixmap.scaled(self.pixmap.size() * self.scale, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        self.image_label.setPixmap(scaled_pixmap)
        self.image_label.adjustSize()
        self.scroll_area.horizontalScrollBar().setValue(self.offset[0])
        self.scroll_area.verticalScrollBar().setValue(self.offset[1])

    def update_zoom(self, value):
        self.scale = value / 100.0
        self.update_image()

    def eventFilter(self, obj, event):
        if obj == self.image_label:
            if event.type() == event.MouseButtonPress and event.button() == Qt.LeftButton:
                self.last_pos = event.pos()
                return True
            elif event.type() == event.MouseMove and event.buttons() == Qt.LeftButton:
                delta = event.pos() - self.last_pos
                self.offset[0] -= delta.x()
                self.offset[1] -= delta.y()
                self.last_pos = event.pos()
                self.update_image()
                return True
            elif event.type() == event.Wheel:
                delta = event.angleDelta().y()
                factor = 1.1 if delta > 0 else 0.909
                new_scale = max(0.5, min(self.scale * factor, 2.0))
                self.zoom_slider.setValue(int(new_scale * 100))
                return True
        return super().eventFilter(obj, event)

class RetailSupportBotApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Retail Support Bot - Knowledge Base Chatbot")
        self.setMinimumSize(800, 500)
        self.is_dark_mode = False

        # Initialize paths
        self.DOCS_DIR = Path("docs")
        self.DATA_DIR = Path("data")
        self.IMAGES_DIR = Path("images")
        self.DATA_DIR.mkdir(exist_ok=True)
        self.IMAGES_DIR.mkdir(exist_ok=True)
        self.DB_PATH = self.DATA_DIR / "knowledge_db.json"
        self.FAQ_DB_PATH = self.DATA_DIR / "support_bot_db.json"
        self.SYNONYMS_PATH = self.DATA_DIR / "synonyms.json"

        # Initialize TinyDB
        self.db = TinyDB(self.DB_PATH)
        self.paragraphs_table = self.db.table('paragraphs')
        self.faq_db = TinyDB(self.FAQ_DB_PATH)
        self.faq_table = self.faq_db.table('faqs')

        # Load synonyms
        try:
            with open(self.SYNONYMS_PATH, 'r') as f:
                self.synonyms = json.load(f)
        except FileNotFoundError:
            self.synonyms = {}
            with open(self.SYNONYMS_PATH, 'w') as f:
                json.dump(self.synonyms, f, indent=4)

        # Chat history
        self.chat_history = []
        self.image_refs = []

        # Central widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)

        # Toolbar
        self.toolbar = self.addToolBar("Actions")
        index_action = QAction("Index Documents", self)
        index_action.triggered.connect(self.batch_index)
        self.toolbar.addAction(index_action)
        upload_action = QAction("Upload Document", self)
        upload_action.triggered.connect(self.upload_document)
        self.toolbar.addAction(upload_action)
        self.theme_action = QAction("Toggle Dark Mode", self)
        self.theme_action.triggered.connect(self.toggle_theme)
        self.toolbar.addAction(self.theme_action)

        # Tabs
        self.tabs = QTabWidget()
        main_layout.addWidget(self.tabs)

        # Initialize tabs
        self.setup_search_tab()
        self.setup_faq_tab()
        self.setup_synonym_tab()
        self.setup_history_tab()
        self.setup_documents_tab()

        # Apply theme
        self.apply_theme()

    def apply_theme(self):
        qss = """
            QWidget { background-color: #f4f4f9; color: #000000; font: 11pt "Helvetica"; }
            QPushButton { background-color: #007bff; color: white; padding: 8px; }
            QPushButton:hover { background-color: #0056b3; }
            QLineEdit, QTextEdit { background-color: #ffffff; color: #000000; font: 10pt "Helvetica"; }
            QComboBox { background-color: #ffffff; color: #000000; }
            QTreeView { background-color: #ffffff; color: #000000; font: 10pt "Helvetica"; }
        """ if not self.is_dark_mode else """
            QWidget { background-color: #2d2d2d; color: #ffffff; font: 11pt "Helvetica"; }
            QPushButton { background-color: #1e90ff; color: #ffffff; padding: 8px; }
            QPushButton:hover { background-color: #4682b4; }
            QLineEdit, QTextEdit { background-color: #3c3c3c; color: #ffffff; font: 10pt "Helvetica"; }
            QComboBox { background-color: #3c3c3c; color: #ffffff; }
            QTreeView { background-color: #3c3c3c; color: #ffffff; font: 10pt "Helvetica"; }
        """
        self.setStyleSheet(qss)
        self.theme_action.setText("Toggle Light Mode" if self.is_dark_mode else "Toggle Dark Mode")

    def toggle_theme(self):
        self.is_dark_mode = not self.is_dark_mode
        self.apply_theme()

    def extract_from_pdf(self, file_path: str, doc_id: str) -> dict:
        try:
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            doc = fitz.open(file_path)
            image_paths = []
            for page_num in range(len(doc)):
                for img in doc[page_num].get_images():
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    if base_image:
                        image_path = self.IMAGES_DIR / f"{doc_id}_img{len(image_paths)}.png"
                        with open(image_path, "wb") as f:
                            f.write(base_image["image"])
                        image_paths.append(str(image_path))
            doc.close()
            return {"text": text, "image_paths": image_paths}
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            QMessageBox.critical(self, "Error", f"Error processing PDF: {str(e)}")
            return {"text": "", "image_paths": []}

    def extract_from_docx(self, file_path: str, doc_id: str) -> dict:
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            image_paths = []
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_bytes = rel.target_part.blob
                    image_ext = rel.target_ref.split('.')[-1]
                    image_path = self.IMAGES_DIR / f"{doc_id}_img{len(image_paths)}.{image_ext}"
                    with open(image_path, "wb") as f:
                        f.write(image_bytes)
                    image_paths.append(str(image_path))
            return {"text": text, "image_paths": image_paths}
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            QMessageBox.critical(self, "Error", f"Error processing DOCX: {str(e)}")
            return {"text": "", "image_paths": []}

    def extract_from_txt(self, file_path: str) -> dict:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return {"text": text, "image_paths": []}
        except Exception as e:
            logger.error(f"Error processing TXT: {str(e)}")
            QMessageBox.critical(self, "Error", f"Error processing TXT: {str(e)}")
            return {"text": "", "image_paths": []}

    def extract_pdf_paragraphs(self, file_path: Path) -> list:
        try:
            with pdfplumber.open(file_path) as pdf:
                paragraphs = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        paras = text.split('\n\n')
                        for para in paras:
                            if para.strip():
                                sentences = sent_tokenize(para.strip())
                                for sentence in sentences:
                                    if len(sentence) > 20:
                                        paragraphs.append(sentence.strip())
                return paragraphs
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            return []

    def extract_docx_paragraphs(self, file_path: Path) -> list:
        try:
            doc = Document(file_path)
            return [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            return []

    def extract_txt_paragraphs(self, file_path: Path) -> list:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return [line.strip() for line in f if line.strip()]
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {str(e)}")
            return []

    def generate_tags(self, text: str, filename: str) -> list:
        stop_words = set(stopwords.words('english'))
        words = word_tokenize(text.lower())
        words = [word for word in words if word.isalpha() and word not in stop_words]
        filename_tags = re.findall(r'\w+', filename.lower())
        word_counts = Counter(words)
        tags = [word for word, _ in word_counts.most_common(5)] + filename_tags
        return list(set(tags))

    def get_all_tags(self):
        tags = set()
        for doc in self.paragraphs_table.all():
            tags.update(doc['tags'])
        return sorted(list(tags))

    def get_indexed_documents(self):
        return sorted(set(doc['filename'] for doc in self.paragraphs_table.all()))

    def setup_search_tab(self):
        chat_widget = QWidget()
        chat_layout = QVBoxLayout(chat_widget)

        # FAQ sidebar
        faq_scroll = QScrollArea()
        faq_scroll.setFixedWidth(200)
        faq_scroll.setWidgetResizable(True)
        faq_widget = QWidget()
        self.faq_layout = QVBoxLayout(faq_widget)
        faq_title = QLabel("Quick FAQ")
        faq_title.setFont(QFont("Helvetica", 12, QFont.Bold))
        self.faq_layout.addWidget(faq_title)
        self.faq_buttons = []
        self.load_faq_buttons()
        faq_scroll.setWidget(faq_widget)

        # Main content
        main_widget = QWidget()
        main_layout = QVBoxLayout(main_widget)

        # Filters
        filter_layout = QHBoxLayout()
        filter_layout.addWidget(QLabel("Filter by Tag:"))
        self.tag_combo = QComboBox()
        self.tag_combo.addItems(self.get_all_tags())
        self.tag_combo.setCurrentIndex(-1)
        filter_layout.addWidget(self.tag_combo)
        filter_layout.addWidget(QLabel("File Type:"))
        self.filetype_combo = QComboBox()
        self.filetype_combo.addItems(["All", "pdf", "docx", "txt"])
        filter_layout.addWidget(self.filetype_combo)
        self.regex_check = QCheckBox("Regex Search")
        filter_layout.addWidget(self.regex_check)
        self.case_sensitive_check = QCheckBox("Case Sensitive")
        filter_layout.addWidget(self.case_sensitive_check)
        self.preset_combo = QComboBox()
        self.preset_combo.addItems(["Custom", "Email", "Phone"])
        self.preset_combo.currentTextChanged.connect(self.set_regex_preset)
        filter_layout.addWidget(self.preset_combo)
        main_layout.addLayout(filter_layout)

        # Results
        self.results_text = QTextEdit()
        self.results_text.setReadOnly(True)
        self.results_text.setFont(QFont("Helvetica", 10))
        self.results_text.setMinimumHeight(150)  # Approx 10 lines
        main_layout.addWidget(self.results_text)

        # Image gallery
        self.image_scroll = QScrollArea()
        self.image_scroll.setFixedHeight(100)
        self.image_scroll.setWidgetResizable(True)
        self.image_widget = QWidget()
        self.image_layout = QHBoxLayout(self.image_widget)
        self.image_scroll.setWidget(self.image_widget)
        main_layout.addWidget(self.image_scroll)

        # Chatbox
        chat_input_layout = QHBoxLayout()
        self.query_input = QLineEdit()
        self.query_input.setFont(QFont("Helvetica", 12))
        self.query_input.setFixedWidth(600)
        self.query_input.returnPressed.connect(self.search)
        chat_input_layout.addWidget(self.query_input)
        send_button = QPushButton("Send")
        send_button.clicked.connect(self.search)
        chat_input_layout.addWidget(send_button)
        clear_button = QPushButton("Clear")
        clear_button.clicked.connect(self.clear_chat)
        chat_input_layout.addWidget(clear_button)
        main_layout.addLayout(chat_input_layout)

        # Combine FAQ and main content
        content_layout = QHBoxLayout()
        content_layout.addWidget(faq_scroll)
        content_layout.addWidget(main_widget)
        chat_layout.addLayout(content_layout)
        self.tabs.addTab(chat_widget, "Chat")

    def setup_faq_tab(self):
        faq_widget = QWidget()
        faq_layout = QVBoxLayout(faq_widget)

        # Input
        input_layout = QHBoxLayout()
        input_layout.addWidget(QLabel("Question:"))
        self.faq_question_input = QLineEdit()
        input_layout.addWidget(self.faq_question_input)
        faq_layout.addLayout(input_layout)
        input_layout = QHBoxLayout()
        input_layout.addWidget(QLabel("Answer:"))
        self.faq_answer_input = QLineEdit()
        input_layout.addWidget(self.faq_answer_input)
        faq_layout.addLayout(input_layout)
        add_button = QPushButton("Add FAQ")
        add_button.clicked.connect(self.add_faq)
        faq_layout.addWidget(add_button)

        # FAQ list
        self.faq_view = QTreeView()
        self.faq_model = QStandardItemModel()
        self.faq_model.setHorizontalHeaderLabels(["Question", "Answer"])
        self.faq_view.setModel(self.faq_model)
        self.faq_view.setColumnWidth(0, 300)
        self.faq_view.setColumnWidth(1, 400)
        self.faq_view.setContextMenuPolicy(Qt.CustomContextMenu)
        self.faq_view.customContextMenuRequested.connect(self.show_faq_context_menu)
        self.faq_view.clicked.connect(self.load_faq)
        faq_layout.addWidget(self.faq_view)
        self.load_faq_list()

        self.tabs.addTab(faq_widget, "FAQ Management")

    def setup_synonym_tab(self):
        synonym_widget = QWidget()
        synonym_layout = QVBoxLayout(synonym_widget)

        # Input
        input_layout = QHBoxLayout()
        input_layout.addWidget(QLabel("Synonym Group (comma-separated):"))
        self.synonym_input = QLineEdit()
        input_layout.addWidget(self.synonym_input)
        synonym_layout.addLayout(input_layout)
        add_button = QPushButton("Add/Edit Synonym Group")
        add_button.clicked.connect(self.add_synonym)
        synonym_layout.addWidget(add_button)

        # Synonym list
        self.synonym_view = QTreeView()
        self.synonym_model = QStandardItemModel()
        self.synonym_model.setHorizontalHeaderLabels(["Group"])
        self.synonym_view.setModel(self.synonym_model)
        self.synonym_view.setColumnWidth(0, 400)
        self.synonym_view.setContextMenuPolicy(Qt.CustomContextMenu)
        self.synonym_view.customContextMenuRequested.connect(self.show_synonym_context_menu)
        self.synonym_view.clicked.connect(self.load_synonym)
        synonym_layout.addWidget(self.synonym_view)
        self.load_synonym_list()

        self.tabs.addTab(synonym_widget, "Synonym Management")

    def setup_history_tab(self):
        history_widget = QWidget()
        history_layout = QVBoxLayout(history_widget)

        history_layout.addWidget(QLabel("Chat History").setFont(QFont("Helvetica", 12, QFont.Bold)))
        self.history_filter = QLineEdit()
        self.history_filter.setPlaceholderText("Filter history...")
        self.history_filter.textChanged.connect(self.filter_history)
        history_layout.addWidget(self.history_filter)
        self.history_view = QTreeView()
        self.history_model = QStandardItemModel()
        self.history_model.setHorizontalHeaderLabels(["Query"])
        self.history_view.setModel(self.history_model)
        self.history_view.setColumnWidth(0, 600)
        self.history_view.doubleClicked.connect(self.load_history_query)
        history_layout.addWidget(self.history_view)
        export_button = QPushButton("Export Chat")
        export_button.clicked.connect(self.export_chat)
        history_layout.addWidget(export_button)
        clear_button = QPushButton("Clear History")
        clear_button.clicked.connect(self.clear_history)
        history_layout.addWidget(clear_button)
        self.load_history_list()

        self.tabs.addTab(history_widget, "Chat History")

    def setup_documents_tab(self):
        documents_widget = QWidget()
        documents_layout = QVBoxLayout(documents_widget)

        documents_layout.addWidget(QLabel("Indexed Documents").setFont(QFont("Helvetica", 12, QFont.Bold)))
        self.documents_view = QTreeView()
        self.documents_model = QStandardItemModel()
        self.documents_model.setHorizontalHeaderLabels(["Filename"])
        self.documents_view.setModel(self.documents_model)
        self.documents_view.setColumnWidth(0, 600)
        documents_layout.addWidget(self.documents_view)
        delete_button = QPushButton("Delete Selected Document")
        delete_button.clicked.connect(self.delete_document)
        documents_layout.addWidget(delete_button)
        self.load_documents_list()

        self.tabs.addTab(documents_widget, "Documents")

    def load_faq_buttons(self):
        # Clear existing buttons
        for btn in self.faq_buttons:
            if btn is not None:
                btn.deleteLater()
        self.faq_buttons = []
        # Load FAQs and create buttons
        for faq in self.faq_table.all():
            if not isinstance(faq, dict) or 'question' not in faq:
                logger.warning(f"Invalid FAQ entry: {faq}")
                continue
            btn = QPushButton(faq['question'])
            if btn is None:
                logger.error("Failed to create QPushButton for FAQ")
                continue
            btn.clicked.connect(lambda checked, q=faq['question']: self.search_faq(q))
            self.faq_layout.addWidget(btn)
            self.faq_buttons.append(btn)
        self.faq_layout.addStretch()

    def load_faq_list(self):
        self.faq_model.removeRows(0, self.faq_model.rowCount())
        for faq in self.faq_table.all():
            if not isinstance(faq, dict) or 'id' not in faq or 'question' not in faq or 'answer' not in faq:
                logger.warning(f"Invalid FAQ entry: {faq}")
                continue
            question_item = QStandardItem(faq['question'])
            answer_item = QStandardItem(faq['answer'])
            question_item.setData(faq['id'], Qt.UserRole)
            self.faq_model.appendRow([question_item, answer_item])

    def load_synonym_list(self):
        self.synonym_model.removeRows(0, self.synonym_model.rowCount())
        for key, words in self.synonyms.items():
            item = QStandardItem(f"{key}: {', '.join(words)}")
            item.setData(key, Qt.UserRole)
            self.synonym_model.appendRow(item)

    def load_history_list(self):
        self.history_model.removeRows(0, self.history_model.rowCount())
        for i, query in enumerate(self.chat_history):
            item = QStandardItem(query)
            item.setData(i, Qt.UserRole)
            self.history_model.appendRow(item)

    def load_documents_list(self):
        self.documents_model.removeRows(0, self.documents_model.rowCount())
        for filename in self.get_indexed_documents():
            self.documents_model.appendRow(QStandardItem(filename))

    def filter_history(self, text):
        for i in range(self.history_model.rowCount()):
            item = self.history_model.item(i)
            item.setHidden(text.lower() not in item.text().lower() if text else False)

    def display_images(self, image_paths):
        # Clear existing image widgets
        for widget in self.image_layout.children():
            if isinstance(widget, QLabel):
                widget.deleteLater()
        self.image_refs = []
        for path in image_paths:
            try:
                pixmap = QPixmap(path)
                if pixmap.isNull():
                    logger.error(f"Failed to load image {path}: Invalid or corrupted image")
                    continue
                scaled_pixmap = pixmap.scaled(80, 80, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                label = QLabel()
                if label is None:
                    logger.error(f"Failed to create QLabel for image {path}")
                    continue
                label.setPixmap(scaled_pixmap)
                label.mousePressEvent = lambda event, p=path: self.open_image(p)
                self.image_layout.addWidget(label)
                self.image_refs.append(pixmap)
            except Exception as e:
                logger.error(f"Error loading image {path}: {str(e)}")
        self.image_layout.addStretch()

    def open_image(self, path):
        dialog = ImageZoomDialog(path, self)
        dialog.exec_()

    def set_regex_preset(self, text):
        if text == "Email":
            self.query_input.setText(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b")
        elif text == "Phone":
            self.query_input.setText(r"\b\d{3}-\d{3}-\d{4}\b")
        else:
            self.query_input.clear()

    def search(self):
        query = self.query_input.text().strip()
        if not query:
            QMessageBox.warning(self, "Warning", "Please enter a question")
            return

        self.chat_history.append(query)
        self.load_history_list()
        self.tabs.setCurrentIndex(0)

        expanded_query = set(query.lower().split())
        for word in query.lower().split():
            for group in self.synonyms.values():
                if word in group:
                    expanded_query.update(group)
        expanded_query = list(expanded_query)

        results = []
        image_paths = set()
        if self.regex_check.isChecked():
            try:
                flags = 0 if self.case_sensitive_check.isChecked() else re.IGNORECASE
                pattern = re.compile(query, flags)
                for doc in self.paragraphs_table.all():
                    if self.filetype_combo.currentText() != "All" and doc['filetype'] != self.filetype_combo.currentText():
                        continue
                    if self.tag_combo.currentText() and self.tag_combo.currentText() not in doc['tags']:
                        continue
                    lines = doc['text'].split('\n')
                    for i, line in enumerate(lines, 1):
                        if pattern.search(line):
                            matched_text = line
                            for match in pattern.finditer(line):
                                matched_text = matched_text[:match.start()] + f"<b>{match.group()}</b>" + matched_text[match.end():]
                            results.append({
                                'filename': doc['filename'],
                                'matched_text': matched_text,
                                'line_number': i,
                                'tags': doc['tags'],
                                'score': 100,
                                'image_paths': doc['image_paths']
                            })
                            image_paths.update(doc['image_paths'])
            except re.error:
                QMessageBox.critical(self, "Error", "Invalid regex pattern")
                return
        else:
            for doc in self.paragraphs_table.all():
                if self.filetype_combo.currentText() != "All" and doc['filetype'] != self.filetype_combo.currentText():
                    continue
                if self.tag_combo.currentText() and self.tag_combo.currentText() not in doc['tags']:
                    continue
                score = max(fuzz.partial_ratio(word.lower(), doc['text'].lower()) for word in expanded_query)
                if score > 70:
                    matched_text = doc['text']
                    for word in expanded_query:
                        matched_text = re.sub(rf'\b({word})\b', r'<b>\1</b>', matched_text, flags=re.IGNORECASE)
                    results.append({
                        'filename': doc['filename'],
                        'matched_text': matched_text,
                        'line_number': None,
                        'tags': doc['tags'],
                        'score': score,
                        'image_paths': doc['image_paths']
                    })
                    image_paths.update(doc['image_paths'])

        results = sorted(results, key=lambda x: x['score'], reverse=True)[:10]
        self.results_text.clear()
        html = ""
        grouped_results = {}
        for result in results:
            filename = result['filename']
            if filename not in grouped_results:
                grouped_results[filename] = []
            grouped_results[filename].append((result['matched_text'], result['score'], result['tags'], result['line_number']))

        for filename, paras in grouped_results.items():
            html += f"<h3>📄 {filename}</h3>"
            for text, score, tags, line_number in paras:
                html += f"<p>Score: {score}% | Tags: {', '.join(tags)}"
                if line_number:
                    html += f" | <i>Line: {line_number}</i>"
                html += f"<br>{text}</p>"
        self.results_text.setHtml(html)
        self.display_images(image_paths)

    def search_faq(self, question: str):
        self.query_input.setText(question)
        self.search()

    def load_faq(self, index):
        if index.isValid():
            faq_id = self.faq_model.item(index.row(), 0).data(Qt.UserRole)
            faq = self.faq_table.get(Query().id == faq_id)
            if faq:
                self.faq_question_input.setText(faq['question'])
                self.faq_answer_input.setText(faq['answer'])
                self.selected_faq_id = faq_id

    def show_faq_context_menu(self, pos):
        index = self.faq_view.indexAt(pos)
        if index.isValid():
            menu = QMenu()
            menu.addAction("Edit", self.edit_faq)
            menu.addAction("Delete", self.delete_faq)
            menu.exec_(self.faq_view.mapToGlobal(pos))

    def show_synonym_context_menu(self, pos):
        index = self.synonym_view.indexAt(pos)
        if index.isValid():
            menu = QMenu()
            menu.addAction("Edit", self.add_synonym)
            menu.addAction("Delete", self.delete_synonym)
            menu.exec_(self.synonym_view.mapToGlobal(pos))

    def load_synonym(self, index):
        if index.isValid():
            key = self.synonym_model.item(index.row()).data(Qt.UserRole)
            self.synonym_input.setText(','.join(self.synonyms[key]))

    def load_history_query(self, index):
        if index.isValid():
            query_index = self.history_model.item(index.row()).data(Qt.UserRole)
            self.query_input.setText(self.chat_history[query_index])
            self.search()

    def clear_chat(self):
        self.query_input.clear()
        self.results_text.clear()
        self.display_images([])

    def clear_history(self):
        self.chat_history = []
        self.load_history_list()
        QMessageBox.information(self, "Success", "Chat history cleared")

    def upload_document(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Upload Document", "", "PDF/DOCX/TXT files (*.pdf *.docx *.txt)")
        if not file_path:
            return

        filename = Path(file_path).name
        if filename in self.get_indexed_documents():
            QMessageBox.warning(self, "Warning", f"Document '{filename}' is already indexed")
            return

        doc_id = str(uuid.uuid4())
        if file_path.lower().endswith('.pdf'):
            extracted_data = self.extract_from_pdf(file_path, doc_id)
            paragraphs = self.extract_pdf_paragraphs(file_path)
            filetype = 'pdf'
        elif file_path.lower().endswith('.docx'):
            extracted_data = self.extract_from_docx(file_path, doc_id)
            paragraphs = self.extract_docx_paragraphs(file_path)
            filetype = 'docx'
        else:
            extracted_data = self.extract_from_txt(file_path)
            paragraphs = self.extract_txt_paragraphs(file_path)
            filetype = 'txt'

        for para in paragraphs:
            tags = self.generate_tags(para, filename)
            self.paragraphs_table.insert({
                'filename': filename,
                'filetype': filetype,
                'text': para,
                'tags': tags,
                'image_paths': extracted_data['image_paths']
            })

        self.tag_combo.clear()
        self.tag_combo.addItems(self.get_all_tags())
        self.load_documents_list()
        QMessageBox.information(self, "Success", "Document indexed successfully")

    def batch_index(self):
        self.DOCS_DIR.mkdir(exist_ok=True)
        doc_files = list(self.DOCS_DIR.glob("*.pdf")) + list(self.DOCS_DIR.glob("*.docx")) + list(self.DOCS_DIR.glob("*.txt"))
        indexed_files = self.get_indexed_documents()
        new_files = [f for f in doc_files if f.name not in indexed_files]

        if not new_files:
            QMessageBox.warning(self, "Warning", "No new .pdf, .docx, or .txt files found in docs/ folder")
            return

        progress = QProgressBar(self)
        progress.setMaximum(len(new_files))
        self.statusBar().showMessage("Indexing documents...")
        self.statusBar().addWidget(progress)

        for i, file_path in enumerate(new_files):
            progress.setValue(i + 1)
            doc_id = str(uuid.uuid4())
            if file_path.suffix.lower() == '.pdf':
                extracted_data = self.extract_from_pdf(file_path, doc_id)
                paragraphs = self.extract_pdf_paragraphs(file_path)
                filetype = 'pdf'
            elif file_path.suffix.lower() == '.docx':
                extracted_data = self.extract_from_docx(file_path, doc_id)
                paragraphs = self.extract_docx_paragraphs(file_path)
                filetype = 'docx'
            else:
                extracted_data = self.extract_from_txt(file_path)
                paragraphs = self.extract_txt_paragraphs(file_path)
                filetype = 'txt'

            for para in paragraphs:
                tags = self.generate_tags(para, file_path.name)
                self.paragraphs_table.insert({
                    'filename': file_path.name,
                    'filetype': filetype,
                    'text': para,
                    'tags': tags,
                    'image_paths': extracted_data['image_paths']
                })

        self.tag_combo.clear()
        self.tag_combo.addItems(self.get_all_tags())
        self.load_documents_list()
        self.statusBar().removeWidget(progress)
        self.statusBar().clearMessage()
        QMessageBox.information(self, "Success", f"Indexed {len(new_files)} new documents")

    def add_faq(self):
        question = self.faq_question_input.text().strip()
        answer = self.faq_answer_input.text().strip()
        if not question or not answer:
            QMessageBox.warning(self, "Warning", "Please enter both question and answer")
            return

        faq_id = str(uuid.uuid4())
        self.faq_table.insert({
            "id": faq_id,
            "question": question,
            "answer": answer
        })
        self.load_faq_buttons()
        self.load_faq_list()
        self.faq_question_input.clear()
        self.faq_answer_input.clear()
        QMessageBox.information(self, "Success", "FAQ added successfully")

    def edit_faq(self):
        index = self.faq_view.currentIndex()
        if not index.isValid():
            QMessageBox.warning(self, "Warning", "Please select an FAQ to edit")
            return

        new_question = self.faq_question_input.text().strip()
        new_answer = self.faq_answer_input.text().strip()
        if not new_question or not new_answer:
            QMessageBox.warning(self, "Warning", "Please enter both question and answer")
            return

        faq_id = self.faq_model.item(index.row(), 0).data(Qt.UserRole)
        self.faq_table.update({
            "question": new_question,
            "answer": new_answer
        }, Query().id == faq_id)
        self.load_faq_buttons()
        self.load_faq_list()
        self.faq_question_input.clear()
        self.faq_answer_input.clear()
        QMessageBox.information(self, "Success", "FAQ updated successfully")

    def delete_faq(self):
        index = self.faq_view.currentIndex()
        if not index.isValid():
            QMessageBox.warning(self, "Warning", "Please select an FAQ to delete")
            return

        faq_id = self.faq_model.item(index.row(), 0).data(Qt.UserRole)
        self.faq_table.remove(Query().id == faq_id)
        self.load_faq_buttons()
        self.load_faq_list()
        QMessageBox.information(self, "Success", "FAQ deleted successfully")

    def add_synonym(self):
        synonym_group = self.synonym_input.text().strip()
        if not synonym_group:
            QMessageBox.warning(self, "Warning", "Please enter a comma-separated synonym group")
            return

        words = [w.strip().lower() for w in synonym_group.split(',')]
        if len(words) < 2:
            QMessageBox.warning(self, "Warning", "Please enter at least two synonyms")
            return

        key = words[0]
        self.synonyms[key] = words
        with open(self.SYNONYMS_PATH, 'w') as f:
            json.dump(self.synonyms, f, indent=4)
        self.load_synonym_list()
        self.synonym_input.clear()
        QMessageBox.information(self, "Success", "Synonym group added/updated successfully")

    def delete_synonym(self):
        index = self.synonym_view.currentIndex()
        if not index.isValid():
            QMessageBox.warning(self, "Warning", "Please select a synonym group to delete")
            return

        key = self.synonym_model.item(index.row()).data(Qt.UserRole)
        del self.synonyms[key]
        with open(self.SYNONYMS_PATH, 'w') as f:
            json.dump(self.synonyms, f, indent=4)
        self.load_synonym_list()
        QMessageBox.information(self, "Success", "Synonym group deleted successfully")

    def delete_document(self):
        index = self.documents_view.currentIndex()
        if not index.isValid():
            QMessageBox.warning(self, "Warning", "Please select a document to delete")
            return

        filename = self.documents_model.item(index.row()).text()
        reply = QMessageBox.question(self, "Confirm", f"Are you sure you want to delete '{filename}' and its associated images?", QMessageBox.Yes | QMessageBox.No)
        if reply == QMessageBox.Yes:
            image_paths = set()
            docs_to_delete = self.paragraphs_table.search(Query().filename == filename)
            for doc in docs_to_delete:
                image_paths.update(doc['image_paths'])
            
            for path in image_paths:
                try:
                    os.remove(path)
                except Exception as e:
                    logger.warning(f"Failed to delete image {path}: {str(e)}")
            
            self.paragraphs_table.remove(Query().filename == filename)
            self.tag_combo.clear()
            self.tag_combo.addItems(self.get_all_tags())
            self.load_documents_list()
            QMessageBox.information(self, "Success", f"Document '{filename}' deleted successfully")

    def export_chat(self):
        with open("chat_transcript.txt", "w") as f:
            for query in self.chat_history:
                f.write(f"Query: {query}\n")
        QMessageBox.information(self, "Success", "Chat transcript exported to chat_transcript.txt")

if __name__ == "__main__":
    from PyQt5.QtWidgets import QApplication
    import sys
    app = QApplication(sys.argv)
    window = RetailSupportBotApp()
    window.show()
    sys.exit(app.exec_())
